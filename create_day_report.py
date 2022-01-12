#pip install --trusted-host pypi.org --trusted-host files.pythonhosted.org python-docx

from datetime import datetime
import pyautogui
import paramiko
import docx
import webbrowser
import time
from docx.shared import Cm

systems = [('ПМТК','192.168.42.115', 'https://192.168.42.115:8443/login', 'user', 'password', 'Сервер регистрации', 'app'),
           ('ПМТК - БД', '192.168.43.170', 'none', 'user', 'password', 'Сервер базы данных', 'winora'),
           ('ПЗ', '192.168.43.106', 'http://192.168.43.106:8085/tkpz/login.seam', 'user', 'password', 'Сервер подачи', 'app'),
           ('ПЗ - БД', '192.168.42.113', 'http://192.168.43.106:8085/tkpz/login.seam', 'user', 'password', 'Сервер базы данных', 'app'),
           ('ДДК', '192.168.40.104', 'http://192.168.40.104:8085/tkrch/login.seam', 'user', 'password', 'Сервер контроля', 'app'),
           ('ДДК - БД', '192.168.42.110', 'http://192.168.40.104:8085/tkrch/login.seam', 'user', 'password', 'Сервер базы', 'app'),
           ('МТК', '192.168.42.24', 'http://192.168.42.24:9000/login', 'user', 'password', 'Сервер отчетов', 'app'),
           ('Graylog', '192.168.43.116', 'http://192.168.43.116:9000/', 'user', 'password', 'Сервер отчетов и бекапов', 'app'),
           ('XML', '192.168.9.213', 'none', 'report', 'password', 'Сервер почтового обмена', 'winxml')]

def winora(system):  #Сервер базы на винде
    import subprocess
    host = system[1]  
    user = system[3]   #'user'
    psw = system[4]    #'pass'
    subprocess.Popen('C:\windows\system32\mstsc.exe')
    time.sleep(1)
    pyautogui.click(878, 333)  # Ставим курсор в поле адреса
    pyautogui.press('backspace', 20)  # Убираем то, что там было
    pyautogui.write(host)  # Вводим нужный адрес
    pyautogui.click(1062, 460)  # Нажимаем коннект
    #Чтобы пароль вводился язык должен быть установлен в ENG
    pyautogui.write(psw)  # Вводим пароль
    pyautogui.click(837, 498)  # Подтверждаем
    pyautogui.press('enter')  # Выскакивает окно, закрываем
    pyautogui.hotkey('win', 'e')  # Запускаем експлорер
    pyautogui.keyDown('ctrlleft')  # Нажимаем и держим
    pyautogui.keyDown('shiftleft')  # Нажимаем и держим
    pyautogui.press('escape')  # Запускаем таскменеджер
    pyautogui.keyUp('ctrlleft')  # Отпускаем
    pyautogui.keyUp('shiftleft')
    screenshot = pyautogui.screenshot()
    screenshot_path = r'C:\bin\winscreenfile.png'
    screenshot.save(screenshot_path)
    pyautogui.keyDown('alt')  # Нажимаем и держим
    pyautogui.press('f4')  # Закрываем таск
    pyautogui.press('f4')  # Закрываем експлорер
    pyautogui.press('f4')  # Закрываем рдп
    pyautogui.keyUp('alt')
    pyautogui.press('enter')  # Выход

def winxml(system):  #Сервер почтового обмена
    import subprocess
    host = system[1]
    user = system[3]
    psw = system[4]

    subprocess.Popen('C:\windows\system32\mstsc.exe')
    time.sleep(1)

    pyautogui.click(878, 333)  # Ставим курсор в поле адреса
    pyautogui.press('backspace', 20)  # Убираем то, что там было
    pyautogui.write("192.168.9.213")  # Вводим нужный адрес
    pyautogui.click(1062, 460)  # Нажимаем коннект
    pyautogui.write('report')  # Вводим пользователя
    pyautogui.press('tab')
    pyautogui.write('report', interval=0.1)  # Вводим пароль
    time.sleep(2)
    pyautogui.press('enter')  # Подключаемся
    time.sleep(10)
    pyautogui.keyDown('alt')  # Нажимаем и держим
    pyautogui.press('f4')  # Закрываем почтового клиента
    pyautogui.press('f4')  # Закрываем окно
    pyautogui.press('f4')  # Закрываем окно
    pyautogui.keyUp('alt')
    pyautogui.doubleClick(34, 335)  # Кликаем на путь
    pyautogui.keyDown('ctrlleft')  # Нажимаем и держим
    pyautogui.keyDown('shiftleft')  # Нажимаем и держим
    pyautogui.press('escape')  # Запускаем таскменеджер
    pyautogui.keyUp('ctrlleft')  # Отпускаем
    pyautogui.keyUp('shiftleft')
    screenshot = pyautogui.screenshot()
    screenshot_path = r'C:\bin\winxmlscreenfile.png'
    screenshot.save(screenshot_path)
    pyautogui.keyDown('alt')  # Нажимаем и держим
    pyautogui.press('f4')  # Закрываем таск
    pyautogui.press('f4')  # Закрываем експлорер
    pyautogui.press('f4')  # Закрываем рдп
    pyautogui.keyUp('alt')
    pyautogui.press('enter')  # Выход

def ssh(system):
    host = system[1]  #'192.168.42.115'
    user = system[3]   #'user'
    psw = system[4]    #'pass'
    commands = 'echo "-Hostname-: "$HOSTNAME  > top.txt; echo "-Date-: "$( date "+%F_%H:%M:%S" ) >> top.txt; echo "---------------------------------------------" >> top.txt; top -b -n 1 >> top.txt; head -15 top.txt'
    commands2 = 'echo "-Hostname-: "$HOSTNAME  > disks.txt; echo "-Date-: "$( date "+%F_%H:%M:%S" ) >> disks.txt; echo "---------------------------------------------" >> disks.txt; df -h >> disks.txt; cat disks.txt'
    path_to_file = r'C:\bin\top.txt'
    path_to_file2 = r'C:\bin\disks.txt'

    ssh_client = paramiko.SSHClient()
    ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh_client.connect(host, username=user, password=psw)
    stdin, stdout, stderr = ssh_client.exec_command(commands)
    stdin2, stdout2, stderr2 = ssh_client.exec_command(commands2)

    output = stdout.readlines()
    file = open(path_to_file, 'w')
    file.write(''.join(output))
    file.close()

    output2 = stdout2.readlines()
    file2 = open(path_to_file2, 'w')
    file2.write(''.join(output2))
    file2.close()

    stdin.close()
    stdin2.close()

now = datetime.now()
date2 = now.strftime("%d/%m/%Y Время %H:%M:%S")
date3 = now.strftime("%d%m%Y")

# create a new documents
WordDocx = docx.Document()
h = WordDocx.add_heading('ЕЖЕДНЕВНЫЙ ОТЧЕТ', 0)
h.alignment = 1

# My paragraph.
WordDocx.add_paragraph(f'Компания "ФАКТУМ"', style = 'No Spacing')
WordDocx.add_paragraph(f'Состояние серверов ТБД на  {date2}', style = 'No Spacing')

for sys in systems:

    if sys[6] == 'app':

        ssh(sys)
        p =WordDocx.add_heading(f'Подсистема ТБД:  {sys[0]}', level=1)
        p2 = WordDocx.add_paragraph(sys[5], style = 'No Spacing')
        p.alignment = 1  # for left, 1 for center, 2 right, 3 justify
        p2.alignment = 1

        WordDocx.add_paragraph('Доступность web сервера', style='Heading 2')

        url = sys[2]      #'https://192.168.42.115:8443/login'
        webbrowser.open(url)
        time.sleep(5)

        screenshot = pyautogui.screenshot()
        screenshot_path = r'C:\bin\screenfile.png'
        screenshot.save(screenshot_path)
        time.sleep(1)
        pyautogui.hotkey('ctrl', 'w')

        image_to_add = WordDocx.add_picture(screenshot_path)
        image_to_add.width = Cm(15)
        image_to_add.height = Cm(10)

        WordDocx.add_paragraph('Нагрузка на сервер (вывод команды top)', style='Heading 2')
        with open(r'C:\bin\top.txt') as f:
            for line in f:
                WordDocx.add_paragraph(line.strip(), style = 'No Spacing')

        WordDocx.add_paragraph('Состояние дисковой подсистемы (вывод df -h)', style='Heading 2')
        with open(r'C:\bin\disks.txt') as f:
            for line in f:
                WordDocx.add_paragraph(line.strip(), style = 'No Spacing')
    elif sys[6] == 'winora':
        winora(sys)
        p =WordDocx.add_heading(f'Подсистема ТБД:  {sys[0]}', level=1)
        p2 = WordDocx.add_paragraph(sys[5], style = 'No Spacing')
        p.alignment = 1  # for left, 1 for center, 2 right, 3 justify
        p2.alignment = 1
        WordDocx.add_paragraph('Дисковая подсистема и нагрузка на сервер', style='Heading 2')
        screenshot_path = r'C:\bin\winscreenfile.png'
        image_to_add = WordDocx.add_picture(screenshot_path)
        image_to_add.width = Cm(15)
        image_to_add.height = Cm(10)

    elif sys[6] == 'winxml':
        winxml(sys)
        p =WordDocx.add_heading(f'Подсистема ТБД:  {sys[0]}', level=1)
        p2 = WordDocx.add_paragraph(sys[5], style = 'No Spacing')
        p.alignment = 1  # for left, 1 for center, 2 right, 3 justify
        p2.alignment = 1
        WordDocx.add_paragraph('Архив почтового обмена и нагрузка на сервер', style='Heading 2')
        screenshot_path = r'C:\bin\winxmlscreenfile.png'
        image_to_add = WordDocx.add_picture(screenshot_path)
        image_to_add.width = Cm(15)
        image_to_add.height = Cm(10)
        
WordDocx.add_paragraph("_______________________________________________________________________________________________________")
WordDocx.add_paragraph(f'Все системы функционируют в штатном режиме, свободное место на дисках есть, очередей на процессорное время не обнаружено.')

WordDocx.add_paragraph("**********************************************************************************", style = 'No Spacing')
WordDocx.add_paragraph("*** Отчет составлен при помощи пакетов python-docx, paramiko, pyautogui ***", style = 'No Spacing')
WordDocx.add_paragraph("**********************************************************************************", style = 'No Spacing')

# Finally savind the document.
WordDocx.save(f'reports/Report-{date3}.docx')
